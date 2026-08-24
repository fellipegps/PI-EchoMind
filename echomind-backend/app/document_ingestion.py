"""Primitivas puras para validacao e extracao documental.

Este modulo nao persiste documentos nem conhece FastAPI, ORM, PGVector ou
chunking. Ele recebe bytes originais e devolve metadados ou blocos de texto.
"""

from __future__ import annotations

import hashlib
import os
import posixpath
import unicodedata
from dataclasses import dataclass
from io import BytesIO
from typing import TYPE_CHECKING, BinaryIO, Iterator

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

if TYPE_CHECKING:
    from sqlalchemy.orm import Session


DEFAULT_MAX_DOCUMENT_SIZE_MB = 10
BYTES_PER_MEGABYTE = 1024 * 1024
READ_CHUNK_SIZE = 64 * 1024
DEFAULT_TEXT_CHUNK_SIZE = 800
DEFAULT_TEXT_CHUNK_OVERLAP = 100
INSTITUTIONAL_CHUNK_SEPARATORS = [
    "\n\n",
    "\n",
    "Art. ",
    "ART. ",
    "§ ",
    "Capitulo ",
    "CAPITULO ",
    "Secao ",
    "SECAO ",
    ". ",
    " ",
    "",
]

ALLOWED_DOCUMENT_MIME_TYPES: dict[str, frozenset[str]] = {
    ".pdf": frozenset({"application/pdf"}),
    ".txt": frozenset({"text/plain"}),
    ".docx": frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    ),
}


class DocumentValidationError(Exception):
    """Erro de dominio previsivel na validacao de um upload documental."""


class InvalidDocumentConfigurationError(DocumentValidationError):
    """A configuracao de tamanho maximo nao e valida."""


class InvalidDocumentFilenameError(DocumentValidationError):
    """O nome do arquivo nao produz um basename seguro."""


class UnsupportedDocumentTypeError(DocumentValidationError):
    """A extensao do arquivo nao pertence a allowlist documental."""


class InvalidDocumentMimeTypeError(DocumentValidationError):
    """O MIME informado nao e compativel com a extensao permitida."""


class EmptyDocumentError(DocumentValidationError):
    """O arquivo nao possui bytes para processamento."""


class DocumentTooLargeError(DocumentValidationError):
    """O arquivo excede o tamanho maximo configurado."""


class DocumentExtractionError(Exception):
    """Erro de dominio previsivel durante a extracao de texto."""


class InvalidDocumentContentError(DocumentExtractionError):
    """Os bytes nao representam um documento extraivel valido."""


class InvalidTextEncodingError(DocumentExtractionError):
    """Os bytes TXT nao usam um encoding permitido."""


class EmptyExtractedDocumentError(DocumentExtractionError):
    """O documento nao contem texto utilizavel."""


class InvalidPdfError(DocumentExtractionError):
    """Os bytes nao representam um PDF que possa ser aberto."""


class EncryptedPdfError(DocumentExtractionError):
    """O PDF esta criptografado e nao pode ser aberto sem senha."""


class PdfOcrNotSupportedError(DocumentExtractionError):
    """O PDF nao possui camada textual suficiente para este MVP."""


class DocumentChunkingError(DocumentExtractionError):
    """A entrada nao pode ser convertida em chunks textuais."""


@dataclass(frozen=True, slots=True)
class ValidatedDocument:
    """Metadados derivados de bytes originais ja validados."""

    filename: str
    mime_type: str
    size_bytes: int
    sha256: str


@dataclass(frozen=True, slots=True)
class ExtractedTextBlock:
    """Unidade de texto ordenada, pronta para um splitter futuro."""

    text: str
    section: str | None = None
    page: int | None = None


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    """Resultado comum de extratores, sem acoplamento a banco ou LangChain."""

    blocks: tuple[ExtractedTextBlock, ...]

    @property
    def text(self) -> str:
        """Texto completo em ordem documental, separado entre blocos."""

        return "\n\n".join(block.text for block in self.blocks)


@dataclass(frozen=True, slots=True)
class ChunkedTextBlock:
    """Chunk puro e rastreavel, sem qualquer acoplamento de persistencia."""

    chunk_index: int
    content: str
    page_start: int | None = None
    page_end: int | None = None
    section_title: str | None = None


@dataclass(frozen=True, slots=True)
class _SourceTextRange:
    start: int
    end: int
    page: int | None
    section: str | None


def get_max_document_size_bytes(value: str | None = None) -> int:
    """Retorna o limite configurado por ``MAX_DOCUMENT_SIZE_MB`` em bytes.

    ``value`` existe para testes e para futuras camadas de configuracao; quando
    omitido, a variavel de ambiente do backend e consultada no momento da
    validacao, sem congelar a configuracao durante a importacao do modulo.
    """

    configured_value = os.getenv("MAX_DOCUMENT_SIZE_MB") if value is None else value
    if configured_value is None or not configured_value.strip():
        megabytes = DEFAULT_MAX_DOCUMENT_SIZE_MB
    else:
        try:
            megabytes = int(configured_value)
        except ValueError as exc:
            raise InvalidDocumentConfigurationError(
                "MAX_DOCUMENT_SIZE_MB deve ser um inteiro positivo."
            ) from exc

    if megabytes <= 0:
        raise InvalidDocumentConfigurationError(
            "MAX_DOCUMENT_SIZE_MB deve ser um inteiro positivo."
        )
    return megabytes * BYTES_PER_MEGABYTE


def sanitize_filename(filename: str) -> str:
    """Normaliza um nome para basename seguro, sem caminhos ou caracteres NUL."""

    if not isinstance(filename, str):
        raise InvalidDocumentFilenameError("O nome do arquivo deve ser texto.")

    normalized = unicodedata.normalize("NFC", filename).replace("\x00", "")
    basename = posixpath.basename(normalized.replace("\\", "/")).strip()
    if basename in {"", ".", ".."}:
        raise InvalidDocumentFilenameError("O nome do arquivo nao e valido.")
    if "/" in basename or "\\" in basename or "\x00" in basename:
        raise InvalidDocumentFilenameError("O nome do arquivo nao e seguro.")
    return basename


def normalize_mime_type(mime_type: str | None) -> str:
    """Remove parametros do Content-Type e devolve somente o media type."""

    if not isinstance(mime_type, str):
        raise InvalidDocumentMimeTypeError("O MIME do arquivo deve ser informado.")
    normalized = mime_type.split(";", maxsplit=1)[0].strip().lower()
    if not normalized:
        raise InvalidDocumentMimeTypeError("O MIME do arquivo deve ser informado.")
    return normalized


def validate_document_metadata(filename: str, mime_type: str | None) -> tuple[str, str]:
    """Valida a allowlist dupla: extensao do basename e MIME compativel."""

    safe_filename = sanitize_filename(filename)
    extension = os.path.splitext(safe_filename)[1].lower()
    allowed_mime_types = ALLOWED_DOCUMENT_MIME_TYPES.get(extension)
    if allowed_mime_types is None:
        raise UnsupportedDocumentTypeError("Extensao documental nao permitida.")

    normalized_mime_type = normalize_mime_type(mime_type)
    if normalized_mime_type not in allowed_mime_types:
        raise InvalidDocumentMimeTypeError(
            "O MIME informado nao corresponde a extensao do documento."
        )
    return safe_filename, normalized_mime_type


def read_limited_bytes(
    source: BinaryIO,
    *,
    max_size_bytes: int,
    chunk_size: int = READ_CHUNK_SIZE,
) -> bytes:
    """Le no maximo o limite configurado mais um byte para detectar excesso."""

    if max_size_bytes <= 0 or chunk_size <= 0:
        raise ValueError("max_size_bytes e chunk_size devem ser positivos.")

    chunks: list[bytes] = []
    total_size = 0
    while total_size <= max_size_bytes:
        remaining = max_size_bytes + 1 - total_size
        chunk = source.read(min(chunk_size, remaining))
        if not chunk:
            break
        if not isinstance(chunk, bytes):
            raise TypeError("A fonte de upload deve fornecer bytes.")
        chunks.append(chunk)
        total_size += len(chunk)

    if total_size > max_size_bytes:
        raise DocumentTooLargeError("O arquivo excede o tamanho maximo permitido.")
    return b"".join(chunks)


def validate_document_bytes(
    *,
    filename: str,
    mime_type: str | None,
    content: bytes,
    max_size_bytes: int | None = None,
) -> ValidatedDocument:
    """Valida bytes originais e calcula seu SHA-256 deterministico."""

    if not isinstance(content, bytes):
        raise TypeError("O conteudo do documento deve ser bytes.")

    safe_filename, normalized_mime_type = validate_document_metadata(filename, mime_type)
    size_bytes = len(content)
    if size_bytes == 0:
        raise EmptyDocumentError("Arquivos vazios nao sao permitidos.")

    effective_max_size = (
        get_max_document_size_bytes() if max_size_bytes is None else max_size_bytes
    )
    if effective_max_size <= 0:
        raise InvalidDocumentConfigurationError(
            "O tamanho maximo do documento deve ser positivo."
        )
    if size_bytes > effective_max_size:
        raise DocumentTooLargeError("O arquivo excede o tamanho maximo permitido.")

    return ValidatedDocument(
        filename=safe_filename,
        mime_type=normalized_mime_type,
        size_bytes=size_bytes,
        sha256=hashlib.sha256(content).hexdigest(),
    )


def validate_document_stream(
    *,
    filename: str,
    mime_type: str | None,
    source: BinaryIO,
    max_size_bytes: int | None = None,
) -> ValidatedDocument:
    """Valida uma fonte binaria lendo-a de forma limitada."""

    effective_max_size = (
        get_max_document_size_bytes() if max_size_bytes is None else max_size_bytes
    )
    content = read_limited_bytes(source, max_size_bytes=effective_max_size)
    return validate_document_bytes(
        filename=filename,
        mime_type=mime_type,
        content=content,
        max_size_bytes=effective_max_size,
    )


def validate_document_for_tenant(
    db: "Session",
    *,
    tenant_id: str,
    filename: str,
    mime_type: str | None,
    content: bytes,
    max_size_bytes: int | None = None,
) -> ValidatedDocument:
    """Valida upload e rejeita SHA-256 ativo apenas dentro do tenant informado."""

    # A dependencia de persistencia fica nesta fronteira para que as funcoes de
    # nome, MIME, tamanho e hash continuem importaveis e puras isoladamente.
    from .document_repository import (
        DuplicateDocumentError,
        find_active_duplicate_document,
    )

    document = validate_document_bytes(
        filename=filename,
        mime_type=mime_type,
        content=content,
        max_size_bytes=max_size_bytes,
    )
    duplicate = find_active_duplicate_document(
        db,
        tenant_id=tenant_id,
        sha256=document.sha256,
    )
    if duplicate is not None:
        raise DuplicateDocumentError("Documento ativo com o mesmo SHA-256 neste tenant.")
    return document


def _normalize_line_endings(text: str) -> str:
    """Converte CRLF/CR em LF sem colapsar quebras de paragrafo internas."""

    return text.replace("\r\n", "\n").replace("\r", "\n")


def _contains_disallowed_control_characters(text: str) -> bool:
    return any(
        ord(character) < 32 and character not in {"\t", "\n", "\r", "\f"}
        for character in text
    )


def _build_extracted_document(
    blocks: list[ExtractedTextBlock],
) -> ExtractedDocument:
    if not blocks:
        raise EmptyExtractedDocumentError("O documento nao contem texto utilizavel.")
    return ExtractedDocument(blocks=tuple(blocks))


def extract_txt(content: bytes) -> ExtractedDocument:
    """Extrai TXT com UTF-8 e fallback limitado a cp1252.

    O fallback cobre arquivos Windows legados. Ele ainda rejeita NUL e controles
    nao textuais para evitar que bytes binarios sejam aceitos silenciosamente.
    """

    if not isinstance(content, bytes):
        raise InvalidDocumentContentError("O conteudo TXT deve ser bytes.")
    if not content:
        raise EmptyExtractedDocumentError("O documento nao contem texto utilizavel.")

    try:
        decoded = content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            decoded = content.decode("cp1252")
        except UnicodeDecodeError as exc:
            raise InvalidTextEncodingError("O TXT nao usa um encoding permitido.") from exc

    if "\x00" in decoded or _contains_disallowed_control_characters(decoded):
        raise InvalidDocumentContentError("O conteudo TXT nao e valido.")

    text = _normalize_line_endings(decoded).strip()
    if not text:
        raise EmptyExtractedDocumentError("O documento nao contem texto utilizavel.")
    return _build_extracted_document([ExtractedTextBlock(text=text)])


def _iter_docx_body_items(document: Document) -> Iterator[Paragraph | Table]:
    """Itera somente os elementos diretos do body, na ordem do XML DOCX."""

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _paragraph_text(paragraph: Paragraph) -> str:
    return _normalize_line_endings(paragraph.text).strip()


def _table_text(table: Table) -> str:
    """Extrai celulas relevantes em ordem, sem repetir celulas mescladas."""

    seen_cells: list[object] = []
    rows: list[str] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            if any(existing is cell._tc for existing in seen_cells):
                continue
            seen_cells.append(cell._tc)
            cell_text = _normalize_line_endings(cell.text).strip()
            if cell_text:
                cells.append(cell_text)
        if cells:
            rows.append("\t".join(cells))
    return "\n".join(rows)


def _is_heading(paragraph: Paragraph) -> bool:
    style = paragraph.style
    return bool(style and style.style_id.startswith("Heading"))


def extract_docx(content: bytes) -> ExtractedDocument:
    """Extrai parágrafos e tabelas DOCX em sua ordem real de documento."""

    if not isinstance(content, bytes):
        raise InvalidDocumentContentError("O conteudo DOCX deve ser bytes.")
    if not content:
        raise EmptyExtractedDocumentError("O documento nao contem texto utilizavel.")

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise InvalidDocumentContentError("O conteudo DOCX nao e valido.") from exc

    blocks: list[ExtractedTextBlock] = []
    current_section: str | None = None
    for item in _iter_docx_body_items(document):
        if isinstance(item, Paragraph):
            text = _paragraph_text(item)
            if not text:
                continue
            if _is_heading(item):
                current_section = text
                blocks.append(ExtractedTextBlock(text=text))
            else:
                blocks.append(ExtractedTextBlock(text=text, section=current_section))
        else:
            text = _table_text(item)
            if text:
                blocks.append(ExtractedTextBlock(text=text, section=current_section))

    return _build_extracted_document(blocks)


def extract_pdf(content: bytes) -> ExtractedDocument:
    """Extrai somente a camada textual de cada pagina PDF, em ordem 1-based."""

    if not isinstance(content, bytes) or not content:
        raise InvalidPdfError("O conteudo PDF nao e valido.")

    try:
        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted:
            raise EncryptedPdfError("O PDF criptografado nao pode ser aberto.")

        blocks: list[ExtractedTextBlock] = []
        for page_number, page in enumerate(reader.pages, start=1):
            extracted_text = page.extract_text()
            text = _normalize_line_endings(extracted_text or "").strip()
            if text:
                blocks.append(ExtractedTextBlock(text=text, page=page_number))
    except EncryptedPdfError:
        raise
    except Exception as exc:
        raise InvalidPdfError("O conteudo PDF nao e valido.") from exc

    if not blocks:
        raise PdfOcrNotSupportedError(
            "PDF sem camada textual; OCR nao suportado no MVP."
        )
    return ExtractedDocument(blocks=tuple(blocks))


def _assemble_chunking_source(
    document: ExtractedDocument,
) -> tuple[str, list[_SourceTextRange]]:
    if not isinstance(document, ExtractedDocument):
        raise DocumentChunkingError("A entrada de chunking deve ser um documento extraido.")

    parts: list[str] = []
    ranges: list[_SourceTextRange] = []
    position = 0
    for block in document.blocks:
        text = _normalize_line_endings(block.text).strip()
        if not text:
            continue
        if parts:
            position += 2  # Separador estavel entre blocos de extracao.
        start = position
        position += len(text)
        parts.append(text)
        ranges.append(
            _SourceTextRange(
                start=start,
                end=position,
                page=block.page,
                section=block.section,
            )
        )

    if not parts:
        raise EmptyExtractedDocumentError("O documento nao contem texto utilizavel.")
    return "\n\n".join(parts), ranges


def _find_chunk_start(
    source: str,
    chunk: str,
    previous_start: int | None,
    previous_end: int | None,
    chunk_overlap: int,
) -> int:
    if previous_start is None or previous_end is None:
        start = source.find(chunk)
    else:
        earliest_expected_start = max(previous_start + 1, previous_end - chunk_overlap)
        start = source.find(chunk, earliest_expected_start)
        if start == -1:
            start = source.find(chunk, previous_start + 1)
    if start == -1:
        raise DocumentChunkingError("Nao foi possivel rastrear a origem de um chunk.")
    return start


def _chunk_metadata(
    ranges: list[_SourceTextRange],
    *,
    start: int,
    end: int,
) -> tuple[int | None, int | None, str | None]:
    contributors = [
        source_range
        for source_range in ranges
        if source_range.start < end and start < source_range.end
    ]
    pages = [source_range.page for source_range in contributors if source_range.page is not None]
    sections = {source_range.section for source_range in contributors if source_range.section}
    return (
        min(pages) if pages else None,
        max(pages) if pages else None,
        next(iter(sections)) if len(sections) == 1 else None,
    )


def chunk_document(
    document: ExtractedDocument,
    *,
    chunk_size: int = DEFAULT_TEXT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_TEXT_CHUNK_OVERLAP,
) -> tuple[ChunkedTextBlock, ...]:
    """Divide texto extraido de forma deterministica, preservando seus metadados."""

    if chunk_size <= 0 or chunk_overlap < 0 or chunk_overlap >= chunk_size:
        raise DocumentChunkingError("A configuracao de tamanho e overlap nao e valida.")

    source, ranges = _assemble_chunking_source(document)
    splitter = RecursiveCharacterTextSplitter(
        separators=INSTITUTIONAL_CHUNK_SEPARATORS,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        is_separator_regex=False,
        keep_separator="end",
    )

    chunks: list[ChunkedTextBlock] = []
    previous_start: int | None = None
    previous_end: int | None = None
    for content in splitter.split_text(source):
        if not content:
            continue
        start = _find_chunk_start(
            source,
            content,
            previous_start,
            previous_end,
            chunk_overlap,
        )
        end = start + len(content)
        page_start, page_end, section_title = _chunk_metadata(
            ranges,
            start=start,
            end=end,
        )
        chunks.append(
            ChunkedTextBlock(
                chunk_index=len(chunks),
                content=content,
                page_start=page_start,
                page_end=page_end,
                section_title=section_title,
            )
        )
        previous_start = start
        previous_end = end

    if not chunks:
        raise EmptyExtractedDocumentError("O documento nao contem texto utilizavel.")
    return tuple(chunks)
