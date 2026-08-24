"""Contratos unitarios da fronteira de validacao documental."""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter


@pytest.fixture()
def ingestion():
    from app import document_ingestion

    return document_ingestion


@pytest.fixture()
def synthetic_docx_bytes():
    """Cria DOCX minimo em memoria, sem dados reais ou acesso de rede."""

    def build(*, paragraphs: list[tuple[str, str | None]], table_rows=None) -> bytes:
        document = Document()
        for text, style in paragraphs:
            document.add_paragraph(text, style=style)
        if table_rows is not None:
            table = document.add_table(rows=0, cols=len(table_rows[0]))
            for row_values in table_rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, row_values, strict=True):
                    cell.text = value
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    return build


@pytest.fixture()
def synthetic_pdf_bytes():
    """Cria PDFs textuais minimos em memoria, sem arquivos ou rede."""

    def build(pages: list[str | None]) -> bytes:
        font_object_number = 3 + (2 * len(pages))
        page_object_numbers = [3 + (2 * index) for index in range(len(pages))]
        objects: list[bytes] = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            (
                b"<< /Type /Pages /Kids ["
                + b" ".join(f"{number} 0 R".encode() for number in page_object_numbers)
                + f"] /Count {len(pages)} >>".encode()
            ),
        ]
        for index, text in enumerate(pages):
            page_number = page_object_numbers[index]
            content_number = page_number + 1
            objects.append(
                (
                    b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                    + f"/Resources << /Font << /F1 {font_object_number} 0 R >> >> ".encode()
                    + f"/Contents {content_number} 0 R >>".encode()
                )
            )
            if text is None:
                stream = b"q\nQ"
            else:
                escaped_text = (
                    text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
                )
                stream = f"BT /F1 12 Tf 72 720 Td ({escaped_text}) Tj ET".encode()
            objects.append(
                f"<< /Length {len(stream)} >>\nstream\n".encode()
                + stream
                + b"\nendstream"
            )
        objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

        output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets = [0]
        for object_number, body in enumerate(objects, start=1):
            offsets.append(len(output))
            output.extend(f"{object_number} 0 obj\n".encode())
            output.extend(body)
            output.extend(b"\nendobj\n")
        xref_offset = len(output)
        output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
        output.extend(b"0000000000 65535 f \n")
        for offset in offsets[1:]:
            output.extend(f"{offset:010d} 00000 n \n".encode())
        output.extend(
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n".encode()
        )
        return bytes(output)

    return build


@pytest.mark.parametrize(
    ("filename", "mime_type"),
    [
        ("regulamento.pdf", "application/pdf"),
        ("notas.TXT", "text/plain; charset=utf-8"),
        (
            "ata.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    ],
)
def test_allowed_extension_and_matching_mime_are_accepted(
    ingestion, filename, mime_type
) -> None:
    document = ingestion.validate_document_bytes(
        filename=filename,
        mime_type=mime_type,
        content=b"conteudo original",
        max_size_bytes=100,
    )

    assert document.filename == filename
    assert document.mime_type == mime_type.split(";", maxsplit=1)[0].lower()


@pytest.mark.parametrize(
    ("filename", "mime_type", "error_name"),
    [
        ("arquivo.exe", "application/octet-stream", "UnsupportedDocumentTypeError"),
        ("arquivo.pdf", "text/plain", "InvalidDocumentMimeTypeError"),
        ("arquivo.txt", "application/pdf", "InvalidDocumentMimeTypeError"),
        ("arquivo.docx", None, "InvalidDocumentMimeTypeError"),
    ],
)
def test_invalid_extension_or_mime_combination_is_rejected(
    ingestion, filename, mime_type, error_name
) -> None:
    with pytest.raises(getattr(ingestion, error_name)):
        ingestion.validate_document_bytes(
            filename=filename,
            mime_type=mime_type,
            content=b"conteudo original",
            max_size_bytes=100,
        )


def test_empty_document_is_rejected(ingestion) -> None:
    with pytest.raises(ingestion.EmptyDocumentError):
        ingestion.validate_document_bytes(
            filename="vazio.txt",
            mime_type="text/plain",
            content=b"",
            max_size_bytes=1,
        )


def test_document_at_limit_is_accepted_and_above_limit_is_rejected(ingestion) -> None:
    accepted = ingestion.validate_document_bytes(
        filename="limite.txt",
        mime_type="text/plain",
        content=b"12345",
        max_size_bytes=5,
    )

    assert accepted.size_bytes == 5
    with pytest.raises(ingestion.DocumentTooLargeError):
        ingestion.validate_document_bytes(
            filename="excesso.txt",
            mime_type="text/plain",
            content=b"123456",
            max_size_bytes=5,
        )


def test_max_document_size_is_configurable(monkeypatch, ingestion) -> None:
    monkeypatch.setenv("MAX_DOCUMENT_SIZE_MB", "2")

    assert ingestion.get_max_document_size_bytes() == 2 * 1024 * 1024
    with pytest.raises(ingestion.InvalidDocumentConfigurationError):
        ingestion.get_max_document_size_bytes("zero")


def test_stream_is_read_only_until_the_first_byte_over_the_limit(ingestion) -> None:
    source = BytesIO(b"123456")

    with pytest.raises(ingestion.DocumentTooLargeError):
        ingestion.read_limited_bytes(source, max_size_bytes=5, chunk_size=2)

    assert source.tell() == 6


def test_sha256_is_deterministic_and_sensitive_to_original_bytes(ingestion) -> None:
    first = ingestion.validate_document_bytes(
        filename="original.txt",
        mime_type="text/plain",
        content=b"linha 1\r\nlinha 2",
        max_size_bytes=100,
    )
    same_bytes = ingestion.validate_document_bytes(
        filename="outro-nome.txt",
        mime_type="text/plain",
        content=b"linha 1\r\nlinha 2",
        max_size_bytes=100,
    )
    changed_bytes = ingestion.validate_document_bytes(
        filename="normalizado.txt",
        mime_type="text/plain",
        content=b"linha 1\nlinha 2",
        max_size_bytes=100,
    )

    assert first.sha256 == same_bytes.sha256
    assert first.sha256 != changed_bytes.sha256


@pytest.mark.parametrize(
    ("raw_filename", "expected_filename"),
    [
        ("../../segredo.txt", "segredo.txt"),
        (r"C:\upload\relatorio.pdf", "relatorio.pdf"),
        ("  pasta/ata\x00.docx  ", "ata.docx"),
    ],
)
def test_filename_is_sanitized_to_a_safe_basename(
    ingestion, raw_filename, expected_filename
) -> None:
    assert ingestion.sanitize_filename(raw_filename) == expected_filename


@pytest.mark.parametrize("filename", ["", ".", "..", "../"])
def test_filename_without_safe_basename_is_rejected(ingestion, filename) -> None:
    with pytest.raises(ingestion.InvalidDocumentFilenameError):
        ingestion.sanitize_filename(filename)


@pytest.mark.parametrize("status", ["pending", "processing", "ready"])
def test_active_duplicate_is_rejected_only_inside_the_same_tenant(
    db, ingestion, status
) -> None:
    from app import document_repository

    content = b"mesmo conteudo"
    sha256 = ingestion.validate_document_bytes(
        filename="original.txt",
        mime_type="text/plain",
        content=content,
        max_size_bytes=100,
    ).sha256
    existing_document = document_repository.create_document(
        db,
        tenant_id="tenant-a",
        data=document_repository.DocumentCreateData(
            filename="original.txt",
            mime_type="text/plain",
            size_bytes=len(content),
            sha256=sha256,
        ),
    )
    if status in {"processing", "ready"}:
        document_repository.transition_document_status(
            db,
            tenant_id="tenant-a",
            document_id=existing_document.id,
            target_status="processing",
        )
    if status == "ready":
        document_repository.transition_document_status(
            db,
            tenant_id="tenant-a",
            document_id=existing_document.id,
            target_status="ready",
        )

    with pytest.raises(document_repository.DuplicateDocumentError):
        ingestion.validate_document_for_tenant(
            db,
            tenant_id="tenant-a",
            filename="copia.txt",
            mime_type="text/plain",
            content=content,
            max_size_bytes=100,
        )

    other_tenant_document = ingestion.validate_document_for_tenant(
        db,
        tenant_id="tenant-b",
        filename="copia.txt",
        mime_type="text/plain",
        content=content,
        max_size_bytes=100,
    )

    assert other_tenant_document.sha256 == sha256


def test_extract_txt_uses_utf8_and_preserves_paragraph_breaks(ingestion) -> None:
    extracted = ingestion.extract_txt("Titulo\r\n\r\nParagrafo dois\rFim".encode())

    assert extracted.blocks == (
        ingestion.ExtractedTextBlock(text="Titulo\n\nParagrafo dois\nFim"),
    )
    assert extracted.text == "Titulo\n\nParagrafo dois\nFim"


def test_extract_txt_uses_documented_cp1252_fallback(ingestion) -> None:
    extracted = ingestion.extract_txt(b"Institui\xe7\xe3o")

    assert extracted.blocks[0].text == "Instituição"


@pytest.mark.parametrize("content", [b"", b" \r\n\t", b"texto\x00binario", b"\x81"])
def test_extract_txt_rejects_empty_or_invalid_content(ingestion, content) -> None:
    expected_error = (
        ingestion.EmptyExtractedDocumentError if content in {b"", b" \r\n\t"}
        else (ingestion.InvalidTextEncodingError if content == b"\x81" else ingestion.InvalidDocumentContentError)
    )

    with pytest.raises(expected_error):
        ingestion.extract_txt(content)


def test_extract_docx_paragraphs_are_ordered_and_headings_define_section(
    ingestion, synthetic_docx_bytes
) -> None:
    content = synthetic_docx_bytes(
        paragraphs=[("Regulamento", "Heading 1"), ("Artigo primeiro", None)],
    )

    extracted = ingestion.extract_docx(content)

    assert extracted.blocks == (
        ingestion.ExtractedTextBlock(text="Regulamento"),
        ingestion.ExtractedTextBlock(text="Artigo primeiro", section="Regulamento"),
    )


def test_extract_docx_table_cells_are_ordered_without_evident_duplicates(
    ingestion, synthetic_docx_bytes
) -> None:
    content = synthetic_docx_bytes(
        paragraphs=[],
        table_rows=[("Campo", "Valor"), ("Campus", "Centro")],
    )

    extracted = ingestion.extract_docx(content)

    assert extracted.blocks == (
        ingestion.ExtractedTextBlock(text="Campo\tValor\nCampus\tCentro"),
    )


def test_extract_docx_does_not_repeat_horizontally_merged_cells(ingestion) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Cabecalho unido"
    table.cell(1, 0).text = "Esquerda"
    table.cell(1, 1).text = "Direita"
    output = BytesIO()
    document.save(output)

    extracted = ingestion.extract_docx(output.getvalue())

    assert extracted.blocks[0].text == "Cabecalho unido\nEsquerda\tDireita"


def test_extract_docx_preserves_order_between_paragraphs_and_table(ingestion) -> None:
    document = Document()
    document.add_paragraph("Antes")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Meio A"
    table.cell(0, 1).text = "Meio B"
    document.add_paragraph("Depois")
    output = BytesIO()
    document.save(output)

    extracted = ingestion.extract_docx(output.getvalue())

    assert [block.text for block in extracted.blocks] == [
        "Antes",
        "Meio A\tMeio B",
        "Depois",
    ]


@pytest.mark.parametrize("content", [b"", b"nao e um docx"])
def test_extract_docx_rejects_empty_or_invalid_content(ingestion, content) -> None:
    with pytest.raises(
        (ingestion.EmptyExtractedDocumentError, ingestion.InvalidDocumentContentError)
    ):
        ingestion.extract_docx(content)


def test_extract_docx_without_text_is_rejected(ingestion, synthetic_docx_bytes) -> None:
    content = synthetic_docx_bytes(paragraphs=[])

    with pytest.raises(ingestion.EmptyExtractedDocumentError):
        ingestion.extract_docx(content)


def test_extract_pdf_returns_textual_page_with_one_based_page_number(
    ingestion, synthetic_pdf_bytes
) -> None:
    extracted = ingestion.extract_pdf(synthetic_pdf_bytes(["Pagina unica"]))

    assert extracted.blocks == (
        ingestion.ExtractedTextBlock(text="Pagina unica", page=1),
    )


def test_extract_pdf_preserves_multiple_page_order_and_empty_page_gap(
    ingestion, synthetic_pdf_bytes
) -> None:
    extracted = ingestion.extract_pdf(synthetic_pdf_bytes(["Primeira", None, "Terceira"]))

    assert [(block.page, block.text) for block in extracted.blocks] == [
        (1, "Primeira"),
        (3, "Terceira"),
    ]


def test_extract_pdf_rejects_corrupted_content(ingestion) -> None:
    with pytest.raises(ingestion.InvalidPdfError):
        ingestion.extract_pdf(b"%PDF-1.4\nconteudo corrompido")


def test_extract_pdf_rejects_encrypted_pdf(ingestion) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("senha-sintetica")
    output = BytesIO()
    writer.write(output)

    with pytest.raises(ingestion.EncryptedPdfError):
        ingestion.extract_pdf(output.getvalue())


def test_extract_pdf_without_text_layer_reports_unsupported_ocr(
    ingestion, synthetic_pdf_bytes
) -> None:
    with pytest.raises(ingestion.PdfOcrNotSupportedError, match="OCR nao suportado"):
        ingestion.extract_pdf(synthetic_pdf_bytes([None, None]))


@pytest.mark.parametrize("length", [1, 800])
def test_chunk_document_keeps_text_at_or_below_default_limit(ingestion, length) -> None:
    document = ingestion.ExtractedDocument(
        blocks=(ingestion.ExtractedTextBlock(text="A" * length),),
    )

    chunks = ingestion.chunk_document(document)

    assert [chunk.content for chunk in chunks] == ["A" * length]
    assert [chunk.chunk_index for chunk in chunks] == [0]


def test_chunk_document_limits_larger_text_and_preserves_overlap(ingestion) -> None:
    document = ingestion.ExtractedDocument(
        blocks=(ingestion.ExtractedTextBlock(text=("A" * 800) + ("B" * 800) + ("C" * 300)),),
    )

    chunks = ingestion.chunk_document(document)

    assert all(0 < len(chunk.content) <= 800 for chunk in chunks)
    assert chunks[0].content[-100:] == chunks[1].content[:100]
    assert chunks[1].content[-100:] == chunks[2].content[:100]
    assert len({chunk.content for chunk in chunks}) == len(chunks)


def test_chunk_document_prioritizes_institutional_separators(ingestion) -> None:
    assert ingestion.INSTITUTIONAL_CHUNK_SEPARATORS.index("Art. ") < (
        ingestion.INSTITUTIONAL_CHUNK_SEPARATORS.index(". ")
    )
    document = ingestion.ExtractedDocument(
        blocks=(ingestion.ExtractedTextBlock(text="Inicio Art. Segundo Terceiro"),),
    )

    chunks = ingestion.chunk_document(document, chunk_size=14, chunk_overlap=0)

    assert [chunk.content for chunk in chunks] == ["Inicio Art.", "Segundo", "Terceiro"]


def test_chunk_document_is_ordered_deterministic_and_preserves_section(ingestion) -> None:
    document = ingestion.ExtractedDocument(
        blocks=(
            ingestion.ExtractedTextBlock(text="Primeiro bloco", section="Regulamento"),
            ingestion.ExtractedTextBlock(text="Segundo bloco", section="Regulamento"),
        ),
    )

    first_run = ingestion.chunk_document(document, chunk_size=15, chunk_overlap=0)
    second_run = ingestion.chunk_document(document, chunk_size=15, chunk_overlap=0)

    assert first_run == second_run
    assert [chunk.chunk_index for chunk in first_run] == list(range(len(first_run)))
    assert [chunk.content for chunk in first_run] == ["Primeiro bloco", "Segundo bloco"]
    assert {chunk.section_title for chunk in first_run} == {"Regulamento"}


def test_chunk_document_tracks_pdf_page_interval_when_chunk_crosses_pages(ingestion) -> None:
    document = ingestion.ExtractedDocument(
        blocks=(
            ingestion.ExtractedTextBlock(text="A" * 350, page=1),
            ingestion.ExtractedTextBlock(text="B" * 400, page=2),
            ingestion.ExtractedTextBlock(text="C" * 100, page=3),
        ),
    )

    chunks = ingestion.chunk_document(document)

    assert [(chunk.page_start, chunk.page_end) for chunk in chunks] == [(1, 2), (3, 3)]
    assert chunks[0].content.startswith("A" * 20)
    assert chunks[0].content.endswith("B" * 20)


def test_chunk_document_rejects_empty_text_without_empty_chunks(ingestion) -> None:
    document = ingestion.ExtractedDocument(
        blocks=(ingestion.ExtractedTextBlock(text=" \n\t "),),
    )

    with pytest.raises(ingestion.EmptyExtractedDocumentError):
        ingestion.chunk_document(document)
