"""Fixtures exclusivas dos testes com PostgreSQL e pgvector reais."""

from __future__ import annotations

import os
from collections.abc import Generator
from hashlib import sha256
from io import BytesIO
from math import sqrt

import pytest
from docx import Document as DocxDocument
from langchain_core.embeddings import Embeddings
from sqlalchemy import create_engine, text
from sqlalchemy.engine import Engine, make_url
from sqlalchemy.pool import NullPool


EXPECTED_DATABASE_NAME = "echomind_integration"
LOCAL_DATABASE_HOSTS = {"127.0.0.1", "localhost", "::1"}


class DeterministicFakeEmbeddings(Embeddings):
    """Embedding 384d local e estavel, sem modelos ou chamadas externas."""

    dimension = 384

    @classmethod
    def _embed(cls, text_value: str) -> list[float]:
        digest = sha256(text_value.encode("utf-8")).digest()
        values = [digest[index % len(digest)] / 127.5 - 1.0 for index in range(cls.dimension)]
        norm = sqrt(sum(value * value for value in values)) or 1.0
        return [value / norm for value in values]

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [self._embed(text_value) for text_value in texts]

    def embed_query(self, text: str) -> list[float]:
        return self._embed(text)


@pytest.fixture(scope="session")
def deterministic_fake_embeddings() -> DeterministicFakeEmbeddings:
    return DeterministicFakeEmbeddings()


@pytest.fixture(scope="session")
def synthetic_txt_bytes() -> bytes:
    """TXT minimo e textual usado no aceite integrado, sem dados reais."""
    return (
        "Regulamento sintetico de matricula.\n\n"
        "O prazo institucional de resposta e de trinta dias."
    ).encode("utf-8")


@pytest.fixture(scope="session")
def synthetic_docx_bytes() -> bytes:
    """DOCX minimo criado em memoria, sem rede ou arquivo persistente."""
    document = DocxDocument()
    document.add_heading("Norma sintetica", level=1)
    document.add_paragraph("O atendimento documental ocorre em horario comercial.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.fixture(scope="session")
def synthetic_pdf_bytes() -> bytes:
    """PDF textual minimo com uma pagina e camada de texto real."""
    text_value = "Calendario sintetico com inscricoes abertas ate setembro."
    escaped_text = (
        text_value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    )
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped_text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        f"<< /Length {len(stream)} >>\nstream\n".encode()
        + stream
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for object_number, value in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{object_number} 0 obj\n".encode())
        output.extend(value)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode()
    )
    return bytes(output)


@pytest.fixture(scope="session")
def integration_database_url() -> str:
    """Aceita somente o banco local e descartavel reservado para integracao."""
    database_url = os.getenv("DATABASE_URL")
    if not database_url:
        pytest.fail("DATABASE_URL deve apontar para o PostgreSQL descartavel de integracao.")

    parsed_url = make_url(database_url)
    if parsed_url.get_backend_name() != "postgresql":
        pytest.fail("A suite integration exige PostgreSQL real; SQLite nao e aceito.")
    if parsed_url.host not in LOCAL_DATABASE_HOSTS:
        pytest.fail("A suite integration recusa bancos remotos, inclusive staging/producao.")
    if parsed_url.database != EXPECTED_DATABASE_NAME:
        pytest.fail(
            f"O banco descartavel deve se chamar {EXPECTED_DATABASE_NAME!r}; "
            f"recebido {parsed_url.database!r}."
        )

    return database_url


@pytest.fixture(scope="session")
def postgres_engine(integration_database_url: str) -> Generator[Engine, None, None]:
    engine = create_engine(integration_database_url, poolclass=NullPool)
    try:
        with engine.connect() as connection:
            connection.execute(text("SELECT 1"))
        yield engine
    finally:
        engine.dispose()
